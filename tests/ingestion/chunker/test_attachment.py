"""첨부 청킹 검증 (feature4-A: docx / xlsx, feature4-B: csv) — chunking-strategy.md §5.

split_attachment는 첨부 파일을 attachment_type별 전략으로 1차 분할하고,
chunk_attachment는 크기 규칙·메타데이터 부착까지 거쳐 Chunk 목록을 산출한다.
docx/xlsx 픽스처는 samples/attachments/ 의 4건을 쓰고, csv는 tmp_path에 생성한다.
"""

import json
from datetime import datetime
from pathlib import Path

import fitz
import openpyxl
import pytest
from docx import Document as DocxDocument

from app.adapters.json_fixture import JsonFixtureSourceAdapter
from app.ingestion.chunker.attachment import (
    build_attachment_metadata,
    chunk_attachment,
    infer_attachment_type,
    split_attachment,
)
from app.ingestion.chunker.base import ChunkDraft
from app.schemas.chunk import Chunk, make_chunk_id
from app.schemas.enums import AttachmentType, ExtractedFormat, SourceType
from app.schemas.page_object import Attachment, PageObject

SAMPLES_DIR = Path(__file__).resolve().parents[3] / "samples"
ATTACHMENTS_DIR = SAMPLES_DIR / "attachments"

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

_PARENT_PAGE = PageObject(
    page_id="CONF-PAGE-1",
    space_key="CLOUD",
    title="EKS 장애 대응 가이드",
    body_html="<h2>본문</h2>",
    version_number=3,
    last_modified="2026-04-22T08:15:00+09:00",
    allowed_groups=["space:CLOUD"],
    allowed_users=["user:taesung"],
    webui_link="/display/CLOUD/eks",
    labels=["eks", "장애대응"],
    ancestors=["Cloud 운영 문서", "EKS 운영"],
)


def _attachment(filename: str, mime_type: str) -> Attachment:
    """samples/attachments/<filename>을 가리키는 Attachment 픽스처.

    ADR-0001 이후 청커는 ``local_path``를 우선 사용하므로 그 값을 채우고,
    ``download_url``은 사용자 노출용 file:// URI로 둔다.
    """
    local = ATTACHMENTS_DIR / filename
    return Attachment(
        attachment_id="CONF-PAGE-1-att-0",
        filename=filename,
        mime_type=mime_type,
        extracted_text="",
        extracted_format=ExtractedFormat.RAW_TEXT,
        download_url=local.as_uri(),
        local_path=str(local),
        parent_page_id="CONF-PAGE-1",
        last_modified=datetime.fromisoformat("2026-04-20T10:00:00+09:00"),
    )


_DOCX_MANUAL = _attachment("EKS_운영_상세_매뉴얼_v2.3.docx", _DOCX_MIME)
_DOCX_ONBOARD = _attachment("신규입사자_온보딩_체크리스트_2026.docx", _DOCX_MIME)
_XLSX_METRICS = _attachment("모니터링_메트릭_정의서_v1.4.xlsx", _XLSX_MIME)
_XLSX_USAGE = _attachment("EKS_노드_월간_사용량_통계_2026Q1.xlsx", _XLSX_MIME)


# --- infer_attachment_type ---


def test_infer_attachment_type_from_mime() -> None:
    assert infer_attachment_type(_DOCX_MANUAL) is AttachmentType.DOCX
    assert infer_attachment_type(_XLSX_METRICS) is AttachmentType.XLSX


def test_infer_attachment_type_falls_back_to_extension() -> None:
    # mime이 generic이어도 확장자로 판별한다
    generic = _attachment("EKS_운영_상세_매뉴얼_v2.3.docx", "application/octet-stream")
    assert infer_attachment_type(generic) is AttachmentType.DOCX


def test_infer_attachment_type_rejects_unknown() -> None:
    unknown = _attachment("memo.hwp", "application/x-hwp")
    with pytest.raises(ValueError, match="첨부 유형"):
        infer_attachment_type(unknown)


# --- docx 1차 분할 (split_attachment) ---


def test_docx_splits_by_heading_hierarchy() -> None:
    drafts = split_attachment(_DOCX_MANUAL, AttachmentType.DOCX)
    assert all(isinstance(d, ChunkDraft) for d in drafts)
    # Heading 1/2/3 각각이 섹션 경계 — 픽스처는 헤딩 44개
    assert len(drafts) == 44
    headers = {d.section_header for d in drafts}
    assert {"0. 개정 이력", "1.1 문서 목적", "4.2 노드 조인 실패"} <= headers
    # 첨부 섹션은 원자성 없음 (2차 재분할·하한선 병합 대상)
    assert all(d.is_atomic is False for d in drafts)


def test_docx_prepends_preamble_to_first_section() -> None:
    drafts = split_attachment(_DOCX_MANUAL, AttachmentType.DOCX)
    first = drafts[0]
    assert first.section_header == "0. 개정 이력"
    # 첫 헤딩 이전 표지 문단(preamble)이 첫 섹션 도입부에 부착된다
    assert first.text.startswith("EKS 운영 상세 매뉴얼")
    assert "주 담당: 최태성, 신유진" in first.text


def test_docx_converts_table_to_markdown() -> None:
    drafts = split_attachment(_DOCX_MANUAL, AttachmentType.DOCX)
    first = drafts[0]
    # '0. 개정 이력' 섹션에 포함된 표가 markdown으로 변환된다
    assert "| 버전 | 일자 | 주요 변경 내용 | 작성자 |" in first.text
    assert "| --- | --- | --- | --- |" in first.text


def test_docx_headingless_falls_back_to_single_draft(tmp_path: Path) -> None:
    # 헤딩이 없는 docx는 단일 draft로 폴백하고 section_header는 파일명을 쓴다
    document = DocxDocument()
    document.add_paragraph("헤딩 없는 첫 문단입니다.")
    document.add_paragraph("헤딩 없는 둘째 문단입니다.")
    path = tmp_path / "plain.docx"
    document.save(path)
    plain = _attachment("plain.docx", _DOCX_MIME)
    plain = plain.model_copy(update={"local_path": str(path), "download_url": path.as_uri()})

    drafts = split_attachment(plain, AttachmentType.DOCX)
    assert len(drafts) == 1
    assert drafts[0].section_header == "plain.docx"
    assert "헤딩 없는 첫 문단입니다." in drafts[0].text
    assert "헤딩 없는 둘째 문단입니다." in drafts[0].text


# --- xlsx 1차 분할 (split_attachment) ---


def test_xlsx_splits_by_sheet_and_serializes_rows() -> None:
    drafts = split_attachment(_XLSX_METRICS, AttachmentType.XLSX)
    headers = {d.section_header for d in drafts}
    # 시트 단위 분할 — 작은 시트는 '[시트명] 행 N~M', 단일 행 oversize 시트는 '[시트명] 행 N'.
    assert "[개정 이력] 행 1~4" in headers
    # 클러스터 메트릭 시트는 단일 행도 800토큰 초과라 행 단위로 분해된다(P2 보완).
    cluster_drafts = [d for d in drafts if d.section_header.startswith("[클러스터 메트릭]")]
    assert cluster_drafts
    # 각 행이 '[<시트명>] <컬럼>: <값> | ...' 형식으로 직렬화된다 (컬럼명 매 행 동봉)
    assert any(
        "[클러스터 메트릭] 메트릭 ID: CL-001 | 메트릭 이름: kubernetes.node.cpu.usage.pct" in d.text
        for d in cluster_drafts
    )


def test_xlsx_omits_empty_cells(tmp_path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "표"
    sheet.append(["이름", "메모", "비고"])
    sheet.append(["alpha", None, "확인"])  # 가운데 셀 비어 있음
    path = tmp_path / "sparse.xlsx"
    workbook.save(path)
    sparse = _attachment("sparse.xlsx", _XLSX_MIME)
    sparse = sparse.model_copy(update={"local_path": str(path), "download_url": path.as_uri()})

    drafts = split_attachment(sparse, AttachmentType.XLSX)
    text = drafts[0].text
    # 빈 셀은 직렬화에서 생략된다
    assert "이름: alpha" in text
    assert "비고: 확인" in text
    assert "메모:" not in text


def test_xlsx_groups_rows_by_50(tmp_path: Path) -> None:
    # 직렬화 토큰이 작은 시트 → 50행 그룹 경계가 그대로 유지된다
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "수치"
    sheet.append(["idx", "value"])
    for i in range(1, 121):  # 데이터 120행
        sheet.append([i, i * 2])
    path = tmp_path / "rows.xlsx"
    workbook.save(path)
    rows_attach = _attachment("rows.xlsx", _XLSX_MIME)
    rows_attach = rows_attach.model_copy(
        update={"local_path": str(path), "download_url": path.as_uri()}
    )

    drafts = split_attachment(rows_attach, AttachmentType.XLSX)
    headers = [d.section_header for d in drafts]
    assert headers == ["[수치] 행 1~50", "[수치] 행 51~100", "[수치] 행 101~120"]


def test_xlsx_oversized_group_shrinks_to_25(tmp_path: Path) -> None:
    # 50행 그룹 직렬화가 800토큰을 초과하면 25행으로 축소 재분할된다
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "긴행"
    sheet.append(["코드", "설명"])
    long_text = "토큰 수를 늘리기 위한 설명 문장"  # 행당 약 25토큰 → 50행 그룹이 800토큰 초과
    for i in range(1, 61):  # 데이터 60행
        sheet.append([f"C-{i}", long_text])
    path = tmp_path / "long.xlsx"
    workbook.save(path)
    long_attach = _attachment("long.xlsx", _XLSX_MIME)
    long_attach = long_attach.model_copy(
        update={"local_path": str(path), "download_url": path.as_uri()}
    )

    drafts = split_attachment(long_attach, AttachmentType.XLSX)
    headers = [d.section_header for d in drafts]
    # 50행 그룹이 25행으로 축소: 1~25 / 26~50 / 51~60
    assert headers == ["[긴행] 행 1~25", "[긴행] 행 26~50", "[긴행] 행 51~60"]


def test_xlsx_single_row_oversize_splits_with_sliding_window(tmp_path: Path) -> None:
    """단일 행이 800토큰을 넘으면 10행 그룹도 축소 불가 → 슬라이딩 윈도우 분할 (P2 보완)."""
    from app.ingestion.chunker.base import MAX_TOKENS
    from app.ingestion.chunker.tokenizer import count_tokens

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "거대행"
    sheet.append(["코드", "설명"])
    long_text = "토큰 " * 1000
    sheet.append(["BIG-1", long_text])
    path = tmp_path / "big_row.xlsx"
    workbook.save(path)
    big = _attachment("big_row.xlsx", _XLSX_MIME)
    big = big.model_copy(update={"local_path": str(path), "download_url": path.as_uri()})

    drafts = split_attachment(big, AttachmentType.XLSX)
    assert len(drafts) >= 2
    headers = [d.section_header for d in drafts]
    assert all("part" in h.lower() for h in headers)
    for draft in drafts:
        assert count_tokens(draft.text) <= MAX_TOKENS


def test_xlsx_datetime_first_row_is_not_misread_as_header(tmp_path: Path) -> None:
    """첫 행에 datetime이 있어도 _cell_to_str의 isoformat 변환에 헤더로 오인되지 않는다."""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "이벤트"
    sheet.append([datetime(2026, 4, 22, 8, 15, 0), "장애 발생"])
    sheet.append([datetime(2026, 4, 22, 8, 20, 0), "복구"])
    path = tmp_path / "events.xlsx"
    workbook.save(path)
    events = _attachment("events.xlsx", _XLSX_MIME)
    events = events.model_copy(update={"local_path": str(path), "download_url": path.as_uri()})

    drafts = split_attachment(events, AttachmentType.XLSX)
    text = drafts[0].text
    # 첫 행은 데이터로 분류되어야 하므로 col_1/col_2 헤더가 부여된다 (ATTACH_NO_HEADER)
    assert "col_1: 2026-04-22T08:15:00" in text
    assert "col_2: 장애 발생" in text


def test_xlsx_synthesizes_header_when_missing(tmp_path: Path) -> None:
    # 첫 행이 데이터(수치)면 헤더 누락으로 보고 col_1, col_2... 를 부여한다 (ATTACH_NO_HEADER)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "헤더없음"
    sheet.append([10, 20, 30])
    sheet.append([40, 50, 60])
    path = tmp_path / "noheader.xlsx"
    workbook.save(path)
    noheader = _attachment("noheader.xlsx", _XLSX_MIME)
    noheader = noheader.model_copy(update={"local_path": str(path), "download_url": path.as_uri()})

    drafts = split_attachment(noheader, AttachmentType.XLSX)
    text = drafts[0].text
    assert "col_1: 10" in text and "col_2: 20" in text and "col_3: 30" in text
    # 첫 행도 데이터로 포함된다
    assert "col_1: 40" in text


# --- 첨부 메타데이터 (build_attachment_metadata) ---


def test_attachment_metadata_inherits_from_parent_page() -> None:
    draft = ChunkDraft(text="[개정 이력] 버전: v1.0", section_header="[개정 이력] 행 1~4")
    meta = build_attachment_metadata(
        _PARENT_PAGE, _XLSX_METRICS, draft, chunk_index=0, attachment_type=AttachmentType.XLSX
    )
    # 첨부 전용 필드
    assert meta.source_type is SourceType.ATTACHMENT
    assert meta.attachment_id == "CONF-PAGE-1-att-0"
    assert meta.attachment_filename == "모니터링_메트릭_정의서_v1.4.xlsx"
    assert meta.attachment_mime == _XLSX_MIME
    assert meta.extracted_format is ExtractedFormat.SHEET_SERIALIZED
    # doc_type 필드는 첨부의 attachment_type 값을 담는다
    assert meta.doc_type == "xlsx"
    # ACL·페이지 메타는 부모 페이지에서 상속
    assert meta.page_id == "CONF-PAGE-1"
    assert meta.allowed_groups == ["space:CLOUD"]
    assert meta.allowed_users == ["user:taesung"]
    assert meta.labels == ["eks", "장애대응"]
    assert meta.space_key == "CLOUD"
    assert meta.webui_link == "/display/CLOUD/eks"
    assert meta.token_count > 0


def test_attachment_chunk_id_uses_attachment_id() -> None:
    draft = ChunkDraft(text="본문", section_header="섹션")
    meta = build_attachment_metadata(
        _PARENT_PAGE, _DOCX_MANUAL, draft, chunk_index=0, attachment_type=AttachmentType.DOCX
    )
    # chunk_id는 attachment_id를 포함한 결정론적 SHA1
    assert meta.chunk_id == make_chunk_id("CONF-PAGE-1", 0, "CONF-PAGE-1-att-0")
    # 본문 청크(attachment_id 없음)와 다른 id
    assert meta.chunk_id != make_chunk_id("CONF-PAGE-1", 0)


def test_attachment_metadata_extracted_format_for_docx() -> None:
    draft = ChunkDraft(text="본문", section_header="섹션")
    meta = build_attachment_metadata(
        _PARENT_PAGE, _DOCX_MANUAL, draft, chunk_index=0, attachment_type=AttachmentType.DOCX
    )
    assert meta.extracted_format is ExtractedFormat.RAW_TEXT
    assert meta.doc_type == "docx"


# --- chunk_attachment 엔트리 ---


def test_chunk_attachment_docx_returns_indexed_chunks() -> None:
    chunks = chunk_attachment(_DOCX_ONBOARD, _PARENT_PAGE, attachment_type=AttachmentType.DOCX)
    assert len(chunks) >= 1
    assert all(isinstance(c, Chunk) for c in chunks)
    assert [c.metadata.chunk_index for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata.source_type is SourceType.ATTACHMENT for c in chunks)
    assert all(c.metadata.extracted_format is ExtractedFormat.RAW_TEXT for c in chunks)
    assert all(c.text.strip() for c in chunks)


def test_chunk_attachment_xlsx_returns_indexed_chunks() -> None:
    chunks = chunk_attachment(_XLSX_USAGE, _PARENT_PAGE, attachment_type=AttachmentType.XLSX)
    assert len(chunks) >= 1
    assert [c.metadata.chunk_index for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata.source_type is SourceType.ATTACHMENT for c in chunks)
    assert all(c.metadata.extracted_format is ExtractedFormat.SHEET_SERIALIZED for c in chunks)
    # 92행 시트가 50→25행 축소 분할되어 여러 청크로 나뉜다
    assert len(chunks) > 4


def test_chunk_attachment_infers_type_when_omitted() -> None:
    # attachment_type 미지정 → mime 기반 추정 (PoC)
    chunks = chunk_attachment(_XLSX_METRICS, _PARENT_PAGE)
    assert len(chunks) >= 1
    assert all(c.metadata.doc_type == "xlsx" for c in chunks)


def test_chunk_attachment_rejects_unknown_type() -> None:
    # pdf/docx/xlsx/csv 4종 외의 알 수 없는 유형 문자열은 _coerce_attachment_type에서 거부
    docx = _attachment("memo.docx", _DOCX_MIME)
    with pytest.raises(ValueError, match="알 수 없는 attachment_type"):
        chunk_attachment(docx, _PARENT_PAGE, attachment_type="hwp")


# --- csv 1차 분할 (split_attachment / chunk_attachment) ---

_CSV_MIME = "text/csv"


def _csv_attachment(tmp_path: Path, filename: str, data: bytes) -> Attachment:
    """tmp_path에 csv 파일을 쓰고 그것을 가리키는 Attachment 픽스처를 만든다."""
    path = tmp_path / filename
    path.write_bytes(data)
    return _attachment(filename, _CSV_MIME).model_copy(
        update={"local_path": str(path), "download_url": path.as_uri()}
    )


def test_infer_attachment_type_csv() -> None:
    assert infer_attachment_type(_attachment("usage.csv", _CSV_MIME)) is AttachmentType.CSV
    # mime이 일반값이어도 확장자로 fallback
    generic = _attachment("usage.csv", "application/octet-stream")
    assert infer_attachment_type(generic) is AttachmentType.CSV


def test_csv_serializes_rows_with_header(tmp_path: Path) -> None:
    data = "메트릭,임계값,설명\nCPU,80,사용률\n메모리,90,여유\n".encode()
    att = _csv_attachment(tmp_path, "metrics.csv", data)
    drafts = split_attachment(att, AttachmentType.CSV)
    assert len(drafts) == 1
    text = drafts[0].text
    # 시트명은 파일 stem, 컬럼명은 매 행에 반복 부착
    assert "[metrics]" in text
    assert "메트릭: CPU" in text
    assert "임계값: 80" in text
    assert "메트릭: 메모리" in text
    # 헤더 행 자체는 데이터로 직렬화되지 않는다
    assert "메트릭: 메트릭" not in text


def test_csv_no_header_falls_back_to_col_n(tmp_path: Path) -> None:
    # 첫 행이 모두 숫자 → 헤더 없음으로 보아 col_1, col_2... 부여
    data = b"1,2,3\n4,5,6\n"
    att = _csv_attachment(tmp_path, "nohdr.csv", data)
    drafts = split_attachment(att, AttachmentType.CSV)
    text = drafts[0].text
    assert "col_1: 1" in text
    assert "col_3: 3" in text
    assert "col_1: 4" in text


def test_csv_omits_empty_cells(tmp_path: Path) -> None:
    data = b"a,b,c\n1,,3\n"
    att = _csv_attachment(tmp_path, "sparse.csv", data)
    drafts = split_attachment(att, AttachmentType.CSV)
    text = drafts[0].text
    assert "a: 1" in text
    assert "c: 3" in text
    assert "b:" not in text  # 빈 셀은 생략


def test_csv_handles_cp949_encoding(tmp_path: Path) -> None:
    # 한글 cp949 인코딩도 깨지지 않고 읽힌다 (인코딩 fallback)
    data = "이름,부서\n홍길동,클라우드\n".encode("cp949")
    att = _csv_attachment(tmp_path, "names.csv", data)
    drafts = split_attachment(att, AttachmentType.CSV)
    text = drafts[0].text
    assert "이름: 홍길동" in text
    assert "부서: 클라우드" in text


def test_csv_handles_utf8_bom(tmp_path: Path) -> None:
    # Excel이 저장하는 utf-8-sig(BOM)에서도 첫 컬럼명이 깨지지 않는다
    data = "메트릭,값\nCPU,80\n".encode("utf-8-sig")
    att = _csv_attachment(tmp_path, "bom.csv", data)
    drafts = split_attachment(att, AttachmentType.CSV)
    assert "메트릭: CPU" in drafts[0].text


def test_csv_groups_rows_by_50(tmp_path: Path) -> None:
    lines = ["col_a,col_b"] + [f"row{i},val{i}" for i in range(120)]
    data = ("\n".join(lines) + "\n").encode()
    att = _csv_attachment(tmp_path, "many.csv", data)
    drafts = split_attachment(att, AttachmentType.CSV)
    # 120 데이터 행 → 50행 그룹으로 최소 3청크 (oversize 시 더 작게 축소)
    assert len(drafts) >= 3
    assert any(d.section_header.startswith("[many] 행 1~") for d in drafts)


def test_chunk_attachment_csv_returns_indexed_chunks(tmp_path: Path) -> None:
    data = "메트릭,값\nCPU,80\n메모리,90\n".encode()
    att = _csv_attachment(tmp_path, "metrics.csv", data)
    chunks = chunk_attachment(att, _PARENT_PAGE, attachment_type=AttachmentType.CSV)
    assert len(chunks) >= 1
    assert [c.metadata.chunk_index for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata.source_type is SourceType.ATTACHMENT for c in chunks)
    assert all(c.metadata.extracted_format is ExtractedFormat.SHEET_SERIALIZED for c in chunks)
    assert all(c.metadata.doc_type == "csv" for c in chunks)
    # ACL은 부모 페이지에서 상속
    assert all(c.metadata.allowed_groups == _PARENT_PAGE.allowed_groups for c in chunks)


def test_chunk_attachment_infers_csv_when_omitted(tmp_path: Path) -> None:
    data = b"a,b\n1,2\n"
    att = _csv_attachment(tmp_path, "infer.csv", data)
    chunks = chunk_attachment(att, _PARENT_PAGE)  # type 미지정 → mime 추정
    assert len(chunks) >= 1
    assert all(c.metadata.doc_type == "csv" for c in chunks)


# --- pdf 1차 분할 (split_attachment / chunk_attachment) ---

_PDF_MIME = "application/pdf"


def _make_pdf(tmp_path: Path, filename: str, lines: list[tuple[str, int, bool]]) -> Attachment:
    """(텍스트, fontsize, bold) 라인 목록으로 PDF를 생성하고 Attachment 픽스처를 만든다."""
    doc = fitz.open()
    page = doc.new_page()
    y = 72.0
    for text, size, bold in lines:
        page.insert_text((72, y), text, fontsize=size, fontname="hebo" if bold else "helv")
        y += size + 10
    path = tmp_path / filename
    doc.save(str(path))
    doc.close()
    return _attachment(filename, _PDF_MIME).model_copy(
        update={"local_path": str(path), "download_url": path.as_uri()}
    )


def test_infer_attachment_type_pdf() -> None:
    assert infer_attachment_type(_attachment("doc.pdf", _PDF_MIME)) is AttachmentType.PDF
    # mime이 일반값이어도 확장자로 fallback
    assert infer_attachment_type(_attachment("doc.pdf", "application/octet-stream")) is (
        AttachmentType.PDF
    )


def test_pdf_splits_by_font_heuristic(tmp_path: Path) -> None:
    att = _make_pdf(
        tmp_path,
        "guide.pdf",
        [
            ("Introduction", 18, True),
            ("This is the body text of the introduction section here.", 11, False),
            ("It continues with more normal body content.", 11, False),
            ("Architecture", 18, True),
            ("The system has three layers described below.", 11, False),
        ],
    )
    drafts = split_attachment(att, AttachmentType.PDF)
    # 큰 폰트/볼드 짧은 행 → 헤딩, section_header는 'p.<페이지>: <제목>'
    assert [d.section_header for d in drafts] == ["p.1: Introduction", "p.1: Architecture"]
    # 제목이 본문 텍스트 앞에 포함되고, 본문 라인이 해당 섹션에 누적된다
    assert drafts[0].text.startswith("Introduction")
    assert "body text of the introduction" in drafts[0].text
    assert "three layers" in drafts[1].text


def test_pdf_headingless_falls_back_to_single_draft(tmp_path: Path) -> None:
    # 균일 폰트(헤딩 미검출) → 단일 draft (chunk_attachment에서 800토큰 슬라이딩 윈도우)
    lines = [
        (f"Uniform body line number {i} with identical font size.", 11, False) for i in range(6)
    ]
    att = _make_pdf(tmp_path, "flat.pdf", lines)
    drafts = split_attachment(att, AttachmentType.PDF)
    assert len(drafts) == 1
    assert drafts[0].section_header == "flat.pdf"
    assert "Uniform body line number 0" in drafts[0].text


def test_pdf_rejects_encrypted(tmp_path: Path) -> None:
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), "secret content", fontsize=11)
    path = tmp_path / "enc.pdf"
    doc.save(str(path), encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw="o", user_pw="u")
    doc.close()
    enc = _attachment("enc.pdf", _PDF_MIME).model_copy(
        update={"local_path": str(path), "download_url": path.as_uri()}
    )
    with pytest.raises(ValueError, match="ATTACH_ENCRYPTED"):
        split_attachment(enc, AttachmentType.PDF)


def test_chunk_attachment_pdf_returns_indexed_chunks(tmp_path: Path) -> None:
    att = _make_pdf(
        tmp_path,
        "guide.pdf",
        [
            ("Overview", 18, True),
            ("Body content for the overview section goes here in detail.", 11, False),
            ("Setup", 18, True),
            ("Steps to set up the system are described in this section.", 11, False),
        ],
    )
    chunks = chunk_attachment(att, _PARENT_PAGE, attachment_type=AttachmentType.PDF)
    assert len(chunks) >= 1
    assert [c.metadata.chunk_index for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata.source_type is SourceType.ATTACHMENT for c in chunks)
    assert all(c.metadata.extracted_format is ExtractedFormat.RAW_TEXT for c in chunks)
    assert all(c.metadata.doc_type == "pdf" for c in chunks)
    # ACL은 부모 페이지에서 상속
    assert all(c.metadata.allowed_groups == _PARENT_PAGE.allowed_groups for c in chunks)


def test_chunk_attachment_infers_pdf_when_omitted(tmp_path: Path) -> None:
    att = _make_pdf(
        tmp_path,
        "infer.pdf",
        [("Title", 18, True), ("Body text here for the document.", 11, False)],
    )
    chunks = chunk_attachment(att, _PARENT_PAGE)  # type 미지정 → mime 추정
    assert len(chunks) >= 1
    assert all(c.metadata.doc_type == "pdf" for c in chunks)


# --- samples/attachments 통합 청킹 ---


def test_samples_attachments_chunk_without_error() -> None:
    adapter = JsonFixtureSourceAdapter(samples_dir=SAMPLES_DIR)
    pairs = [
        (page, attachment) for page in adapter.fetch_pages() for attachment in page.attachments
    ]
    # confluence 샘플의 첨부 4건 (docx 2 + xlsx 2)
    assert len(pairs) == 4

    format_counts = {ExtractedFormat.RAW_TEXT: 0, ExtractedFormat.SHEET_SERIALIZED: 0}
    for page, attachment in pairs:
        chunks = chunk_attachment(attachment, page)
        assert len(chunks) >= 1, f"attachment {attachment.filename} produced no chunks"
        for index, chunk in enumerate(chunks):
            meta = chunk.metadata
            assert chunk.text.strip()
            assert meta.chunk_index == index
            assert meta.section_header, "section_header must not be empty"
            assert meta.source_type is SourceType.ATTACHMENT
            assert meta.attachment_id == attachment.attachment_id
            assert meta.page_id == page.page_id
            assert meta.allowed_groups == page.allowed_groups
            assert meta.token_count > 0
        format_counts[chunks[0].metadata.extracted_format] += 1

    # docx 2건 → raw_text, xlsx 2건 → sheet_serialized
    assert format_counts[ExtractedFormat.RAW_TEXT] == 2
    assert format_counts[ExtractedFormat.SHEET_SERIALIZED] == 2


def test_samples_attachments_pdf_csv_through_adapter(tmp_path: Path) -> None:
    """PDF/CSV 첨부가 JsonFixtureSourceAdapter 전체 경로를 통과하는지 검증한다.

    canonical samples/ 는 docx/xlsx만 가지므로(다른 테스트가 4건을 단언), PDF/CSV의
    어댑터 경유 통합은 tmp_path에 미니 confluence 픽스처(페이지 1건 + 실제 PDF/CSV 파일)를
    만들어 검증한다. fetch_pages → _map_attachments → chunk_attachment → 메타데이터
    무결성(어댑터가 space_key로 합성한 ACL 상속·source_type·extracted_format)을 확인한다.
    canonical 데이터셋을 건드리지 않아 기존 테스트에 영향이 없다.
    """
    attachments_dir = tmp_path / "attachments"
    attachments_dir.mkdir()
    # 실제 CSV (Excel utf-8-sig BOM)
    (attachments_dir / "월간_사용량.csv").write_bytes(
        "메트릭,값\nCPU,80\n메모리,90\n".encode("utf-8-sig")
    )
    # 실제 PDF (헤딩+본문)
    doc = fitz.open()
    page = doc.new_page()
    y = 72.0
    for text, size, bold in [
        ("Overview", 18, True),
        ("This document describes the monthly usage report in detail.", 11, False),
        ("Details", 18, True),
        ("The detailed breakdown is shown across multiple resource pools here.", 11, False),
    ]:
        page.insert_text((72, y), text, fontsize=size, fontname="hebo" if bold else "helv")
        y += size + 10
    doc.save(str(attachments_dir / "운영_리포트.pdf"))
    doc.close()

    fixture = {
        "single_page_responses": [
            {
                "id": "CONF-INT-1",
                "title": "월간 운영 리포트",
                "space": {"key": "CLOUD"},
                "body": {"storage": {"value": "<h2>본문</h2>"}},
                "version": {"when": "2026-04-22T08:15:00.000+0900", "number": 3},
                "_links": {"webui": "/display/CLOUD/report"},
                "metadata": {"labels": {"results": [{"name": "eks"}]}},
                "ancestors": [{"title": "Cloud 운영 문서"}],
                "attachments": [
                    {"filename": "운영_리포트.pdf", "content_type": "application/pdf"},
                    {"filename": "월간_사용량.csv", "content_type": "text/csv"},
                ],
            }
        ]
    }
    (tmp_path / "confluence_sample_data.json").write_text(
        json.dumps(fixture, ensure_ascii=False), encoding="utf-8"
    )

    adapter = JsonFixtureSourceAdapter(
        samples_dir=tmp_path, fixture_files=["confluence_sample_data.json"]
    )
    pairs = [(p, a) for p in adapter.fetch_pages() for a in p.attachments]
    assert len(pairs) == 2

    formats_seen: set[ExtractedFormat] = set()
    for page_obj, attachment in pairs:
        chunks = chunk_attachment(attachment, page_obj)  # type 미지정 → mime 추정
        assert len(chunks) >= 1, f"attachment {attachment.filename} produced no chunks"
        for index, chunk in enumerate(chunks):
            meta = chunk.metadata
            assert chunk.text.strip()
            assert meta.chunk_index == index
            assert meta.section_header, "section_header must not be empty"
            assert meta.source_type is SourceType.ATTACHMENT
            assert meta.attachment_id == attachment.attachment_id
            assert meta.page_id == page_obj.page_id
            # 어댑터가 space_key로 합성한 ACL을 청크가 상속
            assert meta.allowed_groups == page_obj.allowed_groups == ["space:CLOUD"]
            assert meta.token_count > 0
        formats_seen.add(chunks[0].metadata.extracted_format)

    # pdf → raw_text, csv → sheet_serialized 두 경로 모두 어댑터 경유로 동작
    assert formats_seen == {ExtractedFormat.RAW_TEXT, ExtractedFormat.SHEET_SERIALIZED}


def test_resolve_attachment_path_unquotes_file_uri(tmp_path: Path) -> None:
    """local_path 부재 시 file:// URI 의 percent-encoding(한글 파일명)을 복원한다.

    Path.as_uri() 는 비ASCII 를 percent-encode 하므로 unquote 없이는 인코딩된
    리터럴 경로를 열게 된다(배포 전 점검 2026-06-10 — 잠복 결함 수정 회귀).
    """
    from app.ingestion.chunker.attachment import _resolve_attachment_path

    target = tmp_path / "운영가이드.docx"
    target.write_bytes(b"placeholder")
    attachment = Attachment(
        attachment_id="ATT-kr",
        filename="운영가이드.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        extracted_text="",
        extracted_format=ExtractedFormat.RAW_TEXT,
        download_url=target.as_uri(),
        local_path=None,
        parent_page_id="P1",
        last_modified=datetime.fromisoformat("2026-04-22T08:15:00+09:00"),
    )
    assert _resolve_attachment_path(attachment) == str(target)
